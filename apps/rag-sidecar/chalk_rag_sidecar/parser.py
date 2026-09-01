"""Document parsing adapters used before LightRAG ingestion."""

from __future__ import annotations

import io
import json
import os
from dataclasses import dataclass
from pathlib import Path
import shutil
import subprocess
import tempfile
import time
from typing import Any
import zipfile


class ParserError(RuntimeError):
    """Raised when the selected parser is unavailable or fails."""


def _mineru_trust_env() -> bool:
    """Whether MinerU calls inherit HTTP(S)_PROXY environment variables."""
    return os.environ.get("RAG_MINERU_TRUST_ENV", "false").strip().lower() in {
        "1", "true", "yes"
    }


@dataclass(frozen=True)
class ParsedDocument:
    pages: tuple[str, ...]
    engine: str = "text_only"

    @property
    def page_count(self) -> int:
        return len(self.pages)

    def as_index_text(self) -> str:
        paragraphs: list[str] = []
        for page_number, page in enumerate(self.pages, 1):
            page_paragraphs = [part.strip() for part in page.split("\n\n") if part.strip()]
            if not page_paragraphs and page.strip():
                page_paragraphs = [page.strip()]
            for paragraph_number, paragraph in enumerate(page_paragraphs, 1):
                paragraphs.append(
                    f"[CHALK_PAGE:{page_number}][CHALK_PARAGRAPH:{paragraph_number}]\n{paragraph}"
                )
        return "\n\n".join(paragraphs)


def _text_pages(filename: str, content_type: str, content: bytes) -> tuple[str, ...]:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if content_type == "application/pdf" or suffix == "pdf":
        from pypdf import PdfReader

        pages = tuple((page.extract_text() or "") for page in PdfReader(io.BytesIO(content)).pages)
        return pages or ("",)
    if suffix == "docx" or content_type.endswith("wordprocessingml.document"):
        from docx import Document

        document = Document(io.BytesIO(content))
        return ("\n".join(p.text for p in document.paragraphs),)
    try:
        return (content.decode("utf-8"),)
    except UnicodeDecodeError as exc:
        raise ParserError("Only UTF-8 text, PDF, and DOCX documents are supported") from exc


def _markdown_pages(markdown: str) -> tuple[str, ...]:
    pages = tuple(part for part in markdown.split("\f") if part.strip())
    return pages or (markdown,)


def _markitdown(filename: str, content: bytes) -> tuple[str, ...]:
    try:
        from markitdown import MarkItDown
    except ImportError as exc:
        raise ParserError(
            "markitdown is not installed. Install with `pip install -e '.[markitdown]'`."
        ) from exc
    suffix = Path(filename).suffix or ".bin"
    with tempfile.TemporaryDirectory(prefix="chalk-markitdown-") as tmp:
        source = Path(tmp) / f"input{suffix}"
        source.write_bytes(content)
        try:
            result = MarkItDown().convert(str(source))
        except Exception as exc:  # noqa: BLE001
            raise ParserError(f"markitdown failed to convert {filename}: {exc}") from exc
    text = getattr(result, "text_content", None) or getattr(result, "markdown", None) or ""
    return _markdown_pages(str(text))


def _mineru_pages_from_content_list(path: Path) -> tuple[str, ...] | None:
    candidates = list(path.rglob("*_content_list.json")) + list(path.rglob("content_list.json"))
    if not candidates:
        return None
    try:
        payload = json.loads(candidates[0].read_text(encoding="utf-8"))
    except (OSError, UnicodeDecodeError, json.JSONDecodeError):
        return None
    if not isinstance(payload, list):
        return None
    by_page: dict[int, list[str]] = {}
    for block in payload:
        if not isinstance(block, dict):
            continue
        uses_zero_based_page_idx = "page_idx" in block
        raw_page = block.get("page_idx", block.get("page", 0))
        try:
            page_value = int(raw_page)
            page_number = page_value + 1 if uses_zero_based_page_idx else page_value
        except (TypeError, ValueError):
            page_number = 1
        values: list[str] = []
        for key in ("text", "latex", "content", "html", "table_body"):
            value = block.get(key)
            if isinstance(value, str) and value.strip():
                values.append(value.strip())
        if values:
            by_page.setdefault(max(1, page_number), []).append("\n".join(values))
    if not by_page:
        return None
    return tuple("\n\n".join(by_page[index]) for index in sorted(by_page))


def _mineru_local(filename: str, content: bytes) -> tuple[str, ...]:
    cli = os.environ.get("RAG_MINERU_CLI_PATH", "").strip()
    if not cli:
        cli = shutil.which("mineru") or shutil.which("magic-pdf") or ""
    if not cli:
        raise ParserError(
            "MinerU is selected but no local CLI was found. Install `mineru`, configure "
            "RAG_MINERU_CLI_PATH, or set RAG_MINERU_MODE=cloud."
        )
    configured_cli = os.environ.get("RAG_MINERU_CLI_PATH", "").strip()
    if configured_cli and not (Path(cli).is_file() and os.access(cli, os.X_OK)):
        raise ParserError(f"Configured MinerU CLI path is not executable: {cli}")
    if Path(cli).name == "magic-pdf" and Path(filename).suffix.lower() != ".pdf":
        raise ParserError("The legacy magic-pdf CLI only accepts PDF files; install the current mineru CLI for other formats.")
    suffix = Path(filename).suffix or ".bin"
    with tempfile.TemporaryDirectory(prefix="chalk-mineru-") as tmp:
        root = Path(tmp)
        source = root / f"input{suffix}"
        output = root / "output"
        source.write_bytes(content)
        try:
            completed = subprocess.run(
                [cli, "-p", str(source), "-o", str(output)],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=float(os.environ.get("RAG_MINERU_TIMEOUT_SECONDS", "600")),
                check=False,
                shell=False,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            raise ParserError(f"MinerU local parsing failed for {filename}: {exc}") from exc
        if completed.returncode != 0:
            detail = (completed.stderr or completed.stdout or "").strip().splitlines()[-1:]
            raise ParserError(f"MinerU local parsing failed for {filename}: {' '.join(detail)}")
        parsed_root = output / source.stem
        if not parsed_root.is_dir():
            dirs = [item for item in output.glob("**/*") if item.is_dir()] if output.exists() else []
            parsed_root = max(dirs, key=lambda item: item.stat().st_mtime) if dirs else output
        pages = _mineru_pages_from_content_list(parsed_root)
        if pages:
            return pages
        markdown_files = list(parsed_root.rglob("*.md")) if parsed_root.exists() else []
        if not markdown_files:
            raise ParserError(f"MinerU produced no Markdown output for {filename}")
        return _markdown_pages(max(markdown_files, key=lambda item: item.stat().st_size).read_text(encoding="utf-8"))


def _mineru_cloud(filename: str, content: bytes) -> tuple[str, ...]:
    import httpx

    token = os.environ.get("RAG_MINERU_API_TOKEN", "").strip()
    if not token:
        raise ParserError("RAG_MINERU_API_TOKEN is required when RAG_MINERU_MODE=cloud")
    base_url = os.environ.get("RAG_MINERU_API_BASE_URL", "https://mineru.net").rstrip("/")
    headers = {"Authorization": f"Bearer {token}", "Accept": "application/json"}
    timeout = float(os.environ.get("RAG_MINERU_TIMEOUT_SECONDS", "600"))
    body: dict[str, Any] = {
        "files": [{"name": filename, "is_ocr": os.environ.get("RAG_MINERU_OCR", "false").lower() == "true"}],
        "model_version": os.environ.get("RAG_MINERU_MODEL_VERSION", "pipeline"),
        "enable_formula": os.environ.get("RAG_MINERU_ENABLE_FORMULA", "true").lower() == "true",
        "enable_table": os.environ.get("RAG_MINERU_ENABLE_TABLE", "true").lower() == "true",
    }
    language = os.environ.get("RAG_MINERU_LANGUAGE", "").strip()
    if language:
        body["language"] = language
    try:
        with httpx.Client(timeout=timeout, trust_env=_mineru_trust_env()) as client:
            response = client.post(f"{base_url}/api/v4/file-urls/batch", headers=headers, json=body)
            response.raise_for_status()
            data = response.json().get("data") or {}
            batch_id = str(data.get("batch_id") or "")
            upload_urls = data.get("file_urls") or []
            if not batch_id or not upload_urls:
                raise ParserError("MinerU cloud API did not return an upload URL")
            upload = httpx.put(
                str(upload_urls[0]),
                content=content,
                timeout=timeout,
                trust_env=_mineru_trust_env(),
            )
            upload.raise_for_status()
            deadline = time.monotonic() + timeout
            zip_url = ""
            while time.monotonic() < deadline:
                result = client.get(f"{base_url}/api/v4/extract-results/batch/{batch_id}", headers=headers)
                result.raise_for_status()
                entries = (result.json().get("data") or {}).get("extract_result") or []
                entry = next((item for item in entries if item.get("file_name") == filename), None)
                state = str((entry or {}).get("state") or "").lower()
                if state == "done":
                    zip_url = str((entry or {}).get("full_zip_url") or "")
                    break
                if state == "failed":
                    raise ParserError(f"MinerU cloud parsing failed: {(entry or {}).get('err_msg', 'unknown error')}")
                time.sleep(float(os.environ.get("RAG_MINERU_POLL_SECONDS", "4")))
            if not zip_url:
                raise ParserError(f"MinerU cloud parsing timed out for {filename}")
            archive = client.get(zip_url, timeout=timeout)
            archive.raise_for_status()
        with tempfile.TemporaryDirectory(prefix="chalk-mineru-cloud-") as tmp:
            root = Path(tmp)
            with zipfile.ZipFile(io.BytesIO(archive.content)) as bundle:
                for member in bundle.infolist():
                    target = (root / member.filename).resolve()
                    if target != root and root not in target.parents:
                        raise ParserError("MinerU archive contains an unsafe path")
                    if member.is_dir():
                        target.mkdir(parents=True, exist_ok=True)
                    else:
                        target.parent.mkdir(parents=True, exist_ok=True)
                        target.write_bytes(bundle.read(member))
            pages = _mineru_pages_from_content_list(root)
            if pages:
                return pages
            markdown_files = list(root.rglob("*.md"))
            if markdown_files:
                return _markdown_pages(max(markdown_files, key=lambda item: item.stat().st_size).read_text(encoding="utf-8"))
    except ParserError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise ParserError(f"MinerU cloud parsing failed for {filename}: {exc}") from exc
    raise ParserError(f"MinerU cloud produced no Markdown output for {filename}")


def _mineru(filename: str, content: bytes) -> tuple[str, ...]:
    mode = os.environ.get("RAG_MINERU_MODE", "local").strip().lower()
    if mode == "cloud":
        return _mineru_cloud(filename, content)
    if mode != "local":
        raise ParserError("RAG_MINERU_MODE must be either 'local' or 'cloud'")
    return _mineru_local(filename, content)


def parse_document(
    filename: str,
    content_type: str,
    content: bytes,
    *,
    engine: str | None = None,
) -> ParsedDocument:
    """Parse an upload with ``text_only``, ``markitdown`` or ``mineru``."""
    selected = (engine or os.environ.get("RAG_PARSER_ENGINE", "text_only")).strip().lower().replace("-", "_")
    if selected in {"text_only", "text"}:
        return ParsedDocument(_text_pages(filename, content_type, content), "text_only")
    if selected == "markitdown":
        return ParsedDocument(_markitdown(filename, content), "markitdown")
    if selected == "mineru":
        return ParsedDocument(_mineru(filename, content), "mineru")
    raise ParserError(f"Unknown parser engine: {selected!r}. Use text_only, markitdown, or mineru.")


__all__ = ["ParsedDocument", "ParserError", "parse_document"]
